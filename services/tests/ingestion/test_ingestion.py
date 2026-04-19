"""
Tests for ingestion_service — no real external services used.
All Supabase calls are mocked.

Run with:
    pytest tests/ingestion/test_ingestion.py -v
"""
import base64
import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import numpy as np
import pytest
from fastapi.testclient import TestClient
from PIL import Image

from services import ingestion_service


# ---------------------------------------------------------------------------
# Patch Supabase and settings BEFORE importing the app so the module-level
# `supabase = create_client(...)` call never hits the real network.
# ---------------------------------------------------------------------------
with patch('shared.config.settings') as mock_settings, \
     patch('supabase.create_client', return_value=MagicMock()):
    mock_settings.supabase_url = 'http://fake'
    mock_settings.supabase_key = 'fake-key'
    mock_settings.supabase_bucket = 'fake-bucket'
    from services.ingestion_service.app.main import (
        app,
        pdf_to_text,
        image_to_text,
        docx_to_text,
        hybrid_chunking,
    )

client = TestClient(app)


# ---------------------------------------------------------------------------
# Helpers — minimal in-memory file builders
# ---------------------------------------------------------------------------

def make_pdf_bytes(text: str = 'Bonjour le monde') -> bytes:
    """Creates a minimal valid PDF in memory using only stdlib."""
    import struct
    # Minimal single-page PDF with one text stream
    content = (
        b'%PDF-1.4\n'
        b'1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n'
        b'2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n'
        b'3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] '
        b'/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n'
    )
    stream = f'BT /F1 12 Tf 72 720 Td ({text}) Tj ET'.encode()
    content += (
        b'4 0 obj\n<< /Length ' + str(len(stream)).encode() + b' >>\nstream\n'
        + stream + b'\nendstream\nendobj\n'
        b'5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n'
        b'xref\n0 6\n'
        b'0000000000 65535 f \n'
        b'0000000009 00000 n \n'
        b'0000000058 00000 n \n'
        b'0000000115 00000 n \n'
        b'0000000266 00000 n \n'
        b'0000000360 00000 n \n'
        b'trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n441\n%%EOF\n'
    )
    return content


def make_png_bytes() -> bytes:
    """Creates a tiny valid PNG in memory."""
    img = Image.new('RGB', (100, 30), color=(255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def make_docx_bytes(text: str = 'Politique de télétravail') -> bytes:
    """Creates a minimal valid .docx in memory."""
    import docx as python_docx
    doc = python_docx.Document()
    doc.add_paragraph(text)
    doc.add_paragraph('Les employés peuvent travailler 2 jours par semaine.')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# Unit tests — pure functions, no HTTP
# ===========================================================================

class TestHybridChunking:
    def test_returns_chunks_for_normal_text(self):
        text = 'Le télétravail est autorisé. ' * 50
        chunks = hybrid_chunking(text, 'policy.pdf')
        assert len(chunks) > 0

    def test_source_injected_in_every_chunk(self):
        chunks = hybrid_chunking('Some content about remote work rules.', 'policy.pdf')
        assert all('SOURCE: [policy.pdf]' in c.text for c in chunks)

    def test_content_present_in_chunk(self):
        chunks = hybrid_chunking('Remote work is allowed two days a week.', 'policy.pdf')
        assert all('CONTENT:' in c.text for c in chunks)

    def test_metadata_fields(self):
        chunks = hybrid_chunking('Some document text.', 'guide.pdf')
        for chunk in chunks:
            assert chunk.metadata['source'] == 'guide.pdf'
            assert 'section_id' in chunk.metadata
            assert 'chunk_id' in chunk.metadata
            assert 'original_length' in chunk.metadata

    def test_empty_text_returns_no_chunks(self):
        chunks = hybrid_chunking('', 'empty.pdf')
        assert chunks == []

    def test_whitespace_only_returns_no_chunks(self):
        chunks = hybrid_chunking('   \n\n\t  ', 'blank.pdf')
        assert chunks == []

    def test_large_text_produces_multiple_chunks(self):
        text = 'Voici une règle importante. ' * 300
        chunks = hybrid_chunking(text, 'big.pdf')
        assert len(chunks) > 1

    def test_different_sources_produce_different_prefixes(self):
        chunks_a = hybrid_chunking('Same content.', 'doc_a.pdf')
        chunks_b = hybrid_chunking('Same content.', 'doc_b.pdf')
        assert 'doc_a.pdf' in chunks_a[0].text
        assert 'doc_b.pdf' in chunks_b[0].text


class TestDocxToText:
    def test_extracts_paragraphs(self, tmp_path):
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_paragraph('Premier paragraphe')
        doc.add_paragraph('Deuxième paragraphe')
        path = tmp_path / 'test.docx'
        doc.save(path)

        result = docx_to_text(path)
        assert 'Premier paragraphe' in result
        assert 'Deuxième paragraphe' in result

    def test_empty_paragraphs_excluded(self, tmp_path):
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_paragraph('')
        doc.add_paragraph('   ')
        doc.add_paragraph('Contenu réel')
        path = tmp_path / 'test.docx'
        doc.save(path)

        result = docx_to_text(path)
        assert result.strip() == 'Contenu réel'

    def test_returns_string(self, tmp_path):
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_paragraph('Texte')
        path = tmp_path / 'test.docx'
        doc.save(path)

        result = docx_to_text(path)
        assert isinstance(result, str)


class TestImageToText:
    def test_returns_string(self, tmp_path):
        img_path = tmp_path / 'test.png'
        Image.new('RGB', (200, 50), color=(255, 255, 255)).save(img_path)

        mock_reader = MagicMock()
        mock_reader.readtext.return_value = [
            (None, 'Bonjour', 0.99),
            (None, 'monde', 0.95),
        ]

        result = image_to_text(img_path, mock_reader)
        assert result == 'Bonjour monde'

    def test_empty_image_returns_empty_string(self, tmp_path):
        img_path = tmp_path / 'blank.png'
        Image.new('RGB', (100, 30), color=(255, 255, 255)).save(img_path)

        mock_reader = MagicMock()
        mock_reader.readtext.return_value = []

        result = image_to_text(img_path, mock_reader)
        assert result == ''


# ===========================================================================
# Integration tests — HTTP endpoints with mocked Supabase
# ===========================================================================

class TestHealthEndpoint:
    def test_health_returns_200(self):
        response = client.get('/health')
        assert response.status_code == 200


class TestIngestEndpoint:
    def test_ingest_pdf(self):
        pdf_bytes = make_pdf_bytes('Règles de télétravail')
        payload = {
            'filename': 'policy.pdf',
            'content_b64': base64.b64encode(pdf_bytes).decode(),
        }
        # pdf_to_text may return empty for minimal PDFs — we just check the contract
        response = client.post('/ingest', json=payload)
        assert response.status_code == 200
        assert 'chunks' in response.json()

    def test_ingest_docx(self):
        docx_bytes = make_docx_bytes('Politique de congés annuels')
        payload = {
            'filename': 'policy.docx',
            'content_b64': base64.b64encode(docx_bytes).decode(),
        }
        response = client.post('/ingest', json=payload)
        assert response.status_code == 200
        data = response.json()
        assert 'chunks' in data
        assert len(data['chunks']) > 0
        assert all('SOURCE: [policy.docx]' in c['text'] for c in data['chunks'])

    def test_ingest_image(self):
        png_bytes = make_png_bytes()
        payload = {
            'filename': 'screenshot.png',
            'content_b64': base64.b64encode(png_bytes).decode(),
        }
        mock_reader = MagicMock()
        mock_reader.readtext.return_value = [
            (None, 'Texte extrait', 0.98)
        ]
        with patch('ingestion_service.app.main.load_ocr', return_value=mock_reader):
            response = client.post('/ingest', json=payload)
        assert response.status_code == 200
        assert 'chunks' in response.json()

    def test_ingest_unsupported_type_returns_400(self):
        payload = {
            'filename': 'file.xlsx',
            'content_b64': base64.b64encode(b'fake content').decode(),
        }
        response = client.post('/ingest', json=payload)
        assert response.status_code == 400
        assert 'Unsupported file type' in response.json()['detail']

    def test_ingest_invalid_base64_returns_400(self):
        payload = {
            'filename': 'policy.pdf',
            'content_b64': 'NOT_VALID_BASE64!!!',
        }
        response = client.post('/ingest', json=payload)
        assert response.status_code == 400

    def test_ingest_chunks_have_required_metadata(self):
        docx_bytes = make_docx_bytes('Contenu du document de test pour vérifier les métadonnées.')
        payload = {
            'filename': 'test.docx',
            'content_b64': base64.b64encode(docx_bytes).decode(),
        }
        response = client.post('/ingest', json=payload)
        assert response.status_code == 200
        for chunk in response.json()['chunks']:
            assert chunk['metadata']['source'] == 'test.docx'
            assert 'section_id' in chunk['metadata']
            assert 'chunk_id' in chunk['metadata']

